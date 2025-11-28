import uuid
import logging
import os
import io
from typing import Tuple, Optional
from pathlib import Path
from urllib.parse import urlparse, unquote
from fastapi import UploadFile, HTTPException, BackgroundTasks
from sqlalchemy.orm import Session
import fitz
import docx
import boto3
from src.llm.lite_client import lite_client
from src.vector_db.milvus_client import milvus_client
from src.models.document import HRDocument

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# ----------------------------- S3 -----------------------------
S3_BUCKET = os.getenv("S3_BUCKET_NAME")
AWS_REGION = os.getenv("AWS_REGION")

s3 = boto3.client(
    "s3",
    aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
    aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
    region_name=AWS_REGION,
)




# ----------------------------- Chunker -----------------------------
def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200):
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks


class DocumentService:
    """Upload → Extract → Chunk → Embed → Store in Milvus."""

    def __init__(self):
        self.milvus = milvus_client
        self.llm = lite_client
        logger.info("DocumentService initialized (clean + no classifier)")

    # -------------------- Process from S3 URL --------------------
    async def process_document_from_s3(
        self,
        s3_url: str,
        user: any,
        db: Session,
        background_tasks: Optional[BackgroundTasks] = None
    ) -> HRDocument:
        """Process a document that's already in S3"""
        
        # Extract S3 key from URL
        s3_key = self._extract_s3_key_from_url(s3_url)
        
        # Get file metadata from S3
        try:
            response = s3.head_object(Bucket=S3_BUCKET, Key=s3_key)
            file_size = response['ContentLength']
            content_type = response.get('ContentType', 'application/octet-stream')
        except Exception as e:
            logger.error(f"Failed to get file from S3. Bucket: {S3_BUCKET}, Key: {s3_key}, Error: {e}")
            raise HTTPException(400, f"Failed to get file from S3: {e}")
        
        # If content type is generic, detect from filename extension
        if content_type in ['application/octet-stream', 'binary/octet-stream']:
            content_type = self._detect_content_type_from_filename(s3_key)
            logger.info(f"Detected content type from filename: {content_type}")
        
        # Validate content type
        await self._validate_content_type(content_type)
        
        # Generate file_id and filename
        file_id = str(uuid.uuid4())
        filename = s3_key.split('/')[-1]
        
        # Create document record
        document = HRDocument(
            file_id=file_id,
            filename=filename,
            original_filename=filename,
            s3_key=s3_key,
            s3_url=self.generate_presigned_url(s3_key),
            file_size=file_size,
            content_type=content_type,
            uploaded_by_id=user.id,
            processing_status="pending",
        )
        db.add(document)
        db.commit()
        db.refresh(document)
        
        # Download file from S3 for processing
        file_bytes = self._download_from_s3(s3_key)
        
        if background_tasks:
            background_tasks.add_task(
                self._process_background,
                file_id,
                filename,
                content_type,
                file_bytes,
                document.id,
                db
            )
        else:
            await self._process_background(
                file_id,
                filename,
                content_type,
                file_bytes,
                document.id,
                db
            )
        
        return document

    # -------------------- Upload Entry (S3 URL) --------------------
    async def process_document_upload(
        self,
        s3url: str,
        user: any,
        db: Session,
        background_tasks: Optional[BackgroundTasks] = None
    ) -> HRDocument:
        """
        Process a document from S3 URL.
        This is a wrapper that calls process_document_from_s3.
        """
        return await self.process_document_from_s3(
            s3_url=s3url,
            user=user,
            db=db,
            background_tasks=background_tasks
        )

    # -------------------- Upload Entry (File Upload) --------------------
    async def process_file_upload(
        self,
        file: UploadFile,
        user: any,
        db: Session,
        background_tasks: Optional[BackgroundTasks] = None
    ) -> HRDocument:
        """
        Process a direct file upload.
        Upload file to S3, then process it.
        """
        await self._validate_file(file)
        file_id = str(uuid.uuid4())

        s3_key, file_size = await self._upload_to_s3(file, file_id)

        document = await self._create_document_record(
            db=db,
            file=file,
            user=user,
            file_id=file_id,
            s3_key=s3_key,
            file_size=file_size
        )

        # Make a safe copy for background processing
        file.file.seek(0)
        file_bytes = await file.read()

        if background_tasks:
            background_tasks.add_task(
                self._process_background,
                file_id,
                file.filename,
                file.content_type,
                file_bytes,
                document.id,
                db
            )
        else:
            await self._process_background(
                file_id,
                file.filename,
                file.content_type,
                file_bytes,
                document.id,
                db
            )

        return document

    # -------------------- Helper Methods --------------------
    def _detect_content_type_from_filename(self, filename: str) -> str:
        """Detect content type from file extension"""
        extension = filename.lower().split('.')[-1]
        
        content_type_map = {
            'pdf': 'application/pdf',
            'doc': 'application/msword',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'txt': 'text/plain',
        }
        
        detected_type = content_type_map.get(extension, 'application/octet-stream')
        logger.info(f"Detected content type for .{extension}: {detected_type}")
        return detected_type

    def _extract_s3_key_from_url(self, s3_url: str) -> str:
        """
        Extract S3 key from various S3 URL formats.
        
        Supports:
        - s3://bucket-name/path/to/file.pdf
        - https://bucket-name.s3.region.amazonaws.com/path/to/file.pdf
        - https://s3.region.amazonaws.com/bucket-name/path/to/file.pdf
        - https://bucket-name.s3.amazonaws.com/path/to/file.pdf
        - Direct key: path/to/file.pdf
        """
        try:
            # Remove leading/trailing whitespace
            s3_url = s3_url.strip()
            
            # Handle s3:// protocol
            if s3_url.startswith('s3://'):
                # s3://bucket-name/path/to/file.pdf
                parts = s3_url.replace('s3://', '').split('/', 1)
                if len(parts) > 1:
                    # Skip bucket name, return the key
                    return unquote(parts[1])
                else:
                    raise ValueError("Invalid s3:// URL format")
            
            # Handle HTTPS URLs
            elif s3_url.startswith('http://') or s3_url.startswith('https://'):
                parsed = urlparse(s3_url)
                path = parsed.path
                
                # Remove leading slash
                path = path.lstrip('/')
                
                # URL decode the path
                path = unquote(path)
                
                # Check if bucket name is in the hostname (virtual-hosted style)
                # e.g., https://bucket-name.s3.region.amazonaws.com/path/to/file.pdf
                if '.s3.' in parsed.netloc or '.s3-' in parsed.netloc:
                    # Virtual-hosted style: bucket is in hostname, path is the key
                    return path
                
                # Path-style URL (legacy)
                # e.g., https://s3.region.amazonaws.com/bucket-name/path/to/file.pdf
                elif 's3' in parsed.netloc and 'amazonaws.com' in parsed.netloc:
                    # First part of path is bucket name, rest is key
                    parts = path.split('/', 1)
                    if len(parts) > 1:
                        return parts[1]
                    else:
                        return path
                
                # Generic URL - assume entire path is the key
                else:
                    return path
            
            # Assume it's already a key (no protocol)
            else:
                return unquote(s3_url)
                
        except Exception as e:
            logger.error(f"Failed to extract S3 key from URL: {s3_url}, Error: {e}")
            raise HTTPException(400, f"Invalid S3 URL format: {s3_url}")

    def _download_from_s3(self, s3_key: str) -> bytes:
        """Download file from S3 and return bytes"""
        try:
            logger.info(f"Downloading from S3 - Bucket: {S3_BUCKET}, Key: {s3_key}")
            response = s3.get_object(Bucket=S3_BUCKET, Key=s3_key)
            return response['Body'].read()
        except Exception as e:
            logger.error(f"Failed to download from S3. Bucket: {S3_BUCKET}, Key: {s3_key}, Error: {e}")
            raise HTTPException(500, f"Failed to download from S3: {e}")

    # -------------------- Validation --------------------
    async def _validate_file(self, file: UploadFile):
        allowed = [
            "application/pdf",
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/plain",
        ]
        if file.content_type not in allowed:
            raise HTTPException(400, "Unsupported file type.")

    async def _validate_content_type(self, content_type: str):
        allowed = [
            "application/pdf",
            "application/msword",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/plain",
        ]
        if content_type not in allowed:
            raise HTTPException(
                400, 
                f"Unsupported file type: {content_type}. Allowed types: PDF, DOC, DOCX, TXT"
            )

    # -------------------- Extract (In-Memory) --------------------
    def extract_content_from_bytes(self, file_bytes: bytes, content_type: str) -> str:
        try:
            if content_type == "application/pdf":
                # Use BytesIO for in-memory PDF processing
                pdf_stream = io.BytesIO(file_bytes)
                doc = fitz.open(stream=pdf_stream, filetype="pdf")
                text = "\n".join([p.get_text() for p in doc])
                doc.close()
                return text.strip()

            if content_type in (
                "application/msword",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                docx_stream = io.BytesIO(file_bytes)
                d = docx.Document(docx_stream)
                return "\n".join([p.text for p in d.paragraphs]).strip()

            if content_type == "text/plain":
                return file_bytes.decode("utf-8", "ignore").strip()

            return ""

        except Exception as e:
            logger.error(f"Content extraction failed: {e}")
            raise HTTPException(400, f"Content extraction failed: {e}")

    # -------------------- S3 --------------------
    async def _upload_to_s3(self, file: UploadFile, file_id: str) -> Tuple[str, int]:
        file_bytes = await file.read()
        size = len(file_bytes)

        if size > 50 * 1024 * 1024:
            raise HTTPException(400, "File > 50MB")

        ext = (file.filename or "").split(".")[-1]
        s3_key = f"documents/{file_id}.{ext}"

        try:
            s3.put_object(
                Bucket=S3_BUCKET,
                Key=s3_key,
                Body=file_bytes,
                ContentType=file.content_type
            )
            logger.info(f"File uploaded to S3: {s3_key}")
            return s3_key, size
        except Exception as e:
            logger.error(f"S3 upload failed: {e}")
            raise HTTPException(500, f"S3 upload failed: {e}")

    # -------------------- DB --------------------
    async def _create_document_record(
        self, db: Session, file: UploadFile, user, file_id, s3_key, file_size
    ):
        doc = HRDocument(
            file_id=file_id,
            filename=file.filename,
            original_filename=file.filename,
            s3_key=s3_key,
            s3_url=self.generate_presigned_url(s3_key),
            file_size=file_size,
            content_type=file.content_type,
            uploaded_by_id=user.id,
            processing_status="pending",
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
        return doc

    def generate_presigned_url(self, key: str, expires_in: int = 3600):
        try:
            return s3.generate_presigned_url(
                "get_object",
                Params={"Bucket": S3_BUCKET, "Key": key},
                ExpiresIn=expires_in
            )
        except Exception as e:
            logger.error(f"Failed to generate presigned URL: {e}")
            return None

    # -------------------- Background Processing --------------------
    async def _process_background(
        self,
        file_id: str,
        filename: str,
        content_type: str,
        file_bytes: bytes,
        document_id: int,
        db: Session
    ):
        try:
            doc = db.query(HRDocument).filter(HRDocument.id == document_id).first()
            if not doc:
                raise Exception("Document not found in DB")

            doc.processing_status = "processing"
            db.commit()

            # Extract text (fully in-memory)
            content = self.extract_content_from_bytes(file_bytes, content_type)
            if not content:
                raise Exception("Empty text after extraction")

            # Chunk
            chunks = chunk_text(content)

            # Embed
            embeddings = []
            for ch in chunks:
                emb = self.llm.create_embedding(ch)
                embeddings.append(emb)

            # Store in Milvus
            vector_count = await self.milvus.store_document_embeddings(
                file_id=file_id,
                filename=filename,
                content=content,
                embeddings=embeddings,
                category=None,
                chunks=chunks
            )

            doc.vector_count = vector_count
            doc.processing_status = "completed"
            db.commit()
            
            logger.info(f"Document {file_id} processed successfully. Vectors: {vector_count}")

        except Exception as e:
            doc = db.query(HRDocument).filter(HRDocument.id == document_id).first()
            if doc:
                doc.processing_status = "failed"
                doc.error_message = str(e)
                db.commit()
            logger.error(f"Background processing failed for {file_id}: {e}")

    # -------------------- Get & Delete --------------------
    async def get_user_documents(self, user_id: int, db: Session):
        docs = db.query(HRDocument).filter(
            HRDocument.uploaded_by_id == user_id
        ).all()

        for d in docs:
            d.s3_url = self.generate_presigned_url(d.s3_key)

        return docs

    async def delete_document(self, file_id: str, user_id: int, db: Session):
        doc = db.query(HRDocument).filter(
            HRDocument.file_id == file_id,
            HRDocument.uploaded_by_id == user_id
        ).first()

        if not doc:
            raise HTTPException(404, "Document not found")

        # Delete S3
        try:
            s3.delete_object(Bucket=S3_BUCKET, Key=doc.s3_key)
            logger.info(f"Deleted from S3: {doc.s3_key}")
        except Exception as e:
            logger.warning(f"Failed to delete from S3: {e}")

        # Delete from Milvus
        deleted = await self.milvus.delete_by_file_id(file_id)

        db.delete(doc)
        db.commit()

        return {"message": "Document deleted", "vectors_deleted": deleted}


document_service = DocumentService()